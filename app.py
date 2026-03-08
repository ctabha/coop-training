import os
import json
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import pandas as pd
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

STUDENTS_XLSX = os.path.join(DATA_DIR, "students.xlsx")
ASSIGNMENTS_JSON = os.path.join(DATA_DIR, "assignments.json")
LETTER_TEMPLATE = os.path.join(DATA_DIR, "letter_template.docx")

os.makedirs(OUTPUT_DIR, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-key")


def safe_load_json(path, default):
    try:
        if not os.path.exists(path):
            return default
        with open(path, "r", encoding="utf-8") as f:
            txt = f.read().strip()
            if not txt:
                return default
            return json.loads(txt)
    except Exception:
        return default


def safe_save_json(path, data):
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def load_assignments():
    data = safe_load_json(ASSIGNMENTS_JSON, {})
    return data if isinstance(data, dict) else {}


def save_assignments(data):
    safe_save_json(ASSIGNMENTS_JSON, data)


def find_entity_column(df):
    for c in df.columns:
        if str(c).strip() == "جهة التدريب":
            return c
    for c in df.columns:
        if "جهة التدريب" in str(c):
            return c
    return None


def find_name_column(df):
    for c in df.columns:
        if str(c).strip() == "إسم المتدرب":
            return c
    for c in df.columns:
        if str(c).strip() == "اسم المتدرب":
            return c
    for c in df.columns:
        txt = str(c).strip()
        if "اسم المتدرب" in txt or "إسم المتدرب" in txt:
            return c
    return None


def load_students():
    if not os.path.exists(STUDENTS_XLSX):
        raise FileNotFoundError("ملف data/students.xlsx غير موجود")

    df = pd.read_excel(STUDENTS_XLSX)
    df.columns = [str(c).strip() for c in df.columns]

    entity_col = find_entity_column(df)
    name_col = find_name_column(df)

    required_cols = [
        "التخصص",
        "الرقم المرجعي",
        "المدرب",
        "رقم المتدرب",
        "برنامج",
        "رقم الجوال",
    ]

    missing = [c for c in required_cols if c not in df.columns]
    if entity_col is None:
        missing.append("جهة التدريب")
    if name_col is None:
        missing.append("اسم المتدرب")

    if missing:
        raise ValueError("الأعمدة الناقصة في ملف الإكسل: " + " - ".join(missing))

    students = []

    for _, row in df.iterrows():
        trainee_id = str(row["رقم المتدرب"]).strip() if pd.notna(row["رقم المتدرب"]) else ""
        if trainee_id == "" or trainee_id.lower() == "nan":
            continue

        def getv(col):
            if col is None:
                return ""
            val = row[col]
            if pd.isna(val):
                return ""
            s = str(val).strip()
            return "" if s.lower() == "nan" else s

        students.append({
            "trainee_id": trainee_id,
            "trainee_name": getv(name_col),
            "phone": getv("رقم الجوال"),
            "specialization": getv("التخصص"),
            "program": getv("برنامج"),
            "training_entity": getv(entity_col),
            "course_ref": getv("الرقم المرجعي"),
            "college_supervisor": getv("المدرب"),
            "start_date": "",
            "end_date": "",
        })

    return students


def find_student(trainee_id):
    for s in load_students():
        if s["trainee_id"] == str(trainee_id).strip():
            return s
    return None


def compute_slots_from_excel(students):
    slots = {}
    for s in students:
        spec = (s.get("specialization") or "").strip()
        entity = (s.get("training_entity") or "").strip()

        if not spec or not entity:
            continue

        slots.setdefault(spec, {})
        slots[spec][entity] = slots[spec].get(entity, 0) + 1

    return slots


def compute_used_from_assignments(assignments, students):
    id_to_student = {s["trainee_id"]: s for s in students}
    used = {}

    for trainee_id, rec in assignments.items():
        if not isinstance(rec, dict):
            continue

        entity = (rec.get("entity") or "").strip()
        if not entity:
            continue

        st = id_to_student.get(str(trainee_id))
        if not st:
            continue

        spec = (st.get("specialization") or "").strip()
        if not spec:
            continue

        used.setdefault(spec, {})
        used[spec][entity] = used[spec].get(entity, 0) + 1

    return used


def remaining_for_student(student, students, assignments):
    spec = (student.get("specialization") or "").strip()

    slots = compute_slots_from_excel(students)
    used = compute_used_from_assignments(assignments, students)

    spec_slots = slots.get(spec, {})
    spec_used = used.get(spec, {})

    items = []
    total_slots = 0
    total_remaining = 0

    for entity, total in spec_slots.items():
        total_int = int(total)
        used_int = int(spec_used.get(entity, 0))
        remaining = total_int - used_int
        if remaining < 0:
            remaining = 0

        total_slots += total_int
        total_remaining += remaining

        items.append({
            "name": entity,
            "total": total_int,
            "used": used_int,
            "remaining": remaining
        })

    items.sort(key=lambda x: (-x["remaining"], x["name"]))
    available = [x for x in items if x["remaining"] > 0]

    return {
        "all_items": items,
        "available": available,
        "total_slots": total_slots,
        "total_remaining": total_remaining
    }


def replace_in_paragraph(paragraph, mapping):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    for k, v in mapping.items():
        new_text = new_text.replace(k, v)

    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)
            for nested in cell.tables:
                replace_in_table(nested, mapping)


def fill_letter_docx(student, entity_name):
    if not os.path.exists(LETTER_TEMPLATE):
        raise FileNotFoundError("ملف data/letter_template.docx غير موجود")

    doc = Document(LETTER_TEMPLATE)

    mapping = {
        "{{phone}}": student.get("phone", ""),
        "{{trainee_name}}": student.get("trainee_name", ""),
        "{{trainee_id}}": student.get("trainee_id", ""),
        "{{course_ref}}": student.get("course_ref", ""),
        "{{college_supervisor}}": student.get("college_supervisor", ""),
        "{{training_entity}}": entity_name,
        "{{start_date}}": student.get("start_date", ""),
        "{{end_date}}": student.get("end_date", ""),
    }

    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    for t in doc.tables:
        replace_in_table(t, mapping)

    out_name = f"خطاب_توجيه_{student.get('trainee_id','')}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_path


@app.route("/", methods=["GET", "POST", "HEAD"])
def index():
    if request.method == "HEAD":
        return ("", 200)

    if request.method == "GET":
        return render_template("index.html")

    trainee_id = (request.form.get("trainee_id") or "").strip()
    phone_last4 = (request.form.get("phone_last4") or "").strip()

    student = find_student(trainee_id)
    if not student:
        flash("رقم المتدرب غير موجود.", "error")
        return redirect(url_for("index"))

    phone = (student.get("phone") or "").strip()
    if phone_last4:
        if len(phone) < 4 or not phone.endswith(phone_last4):
            flash("آخر 4 أرقام من الجوال غير صحيحة.", "error")
            return redirect(url_for("index"))

    return redirect(url_for("dashboard", trainee_id=trainee_id))


@app.route("/dashboard/<trainee_id>", methods=["GET"])
def dashboard(trainee_id):
    students_all = load_students()
    student = next((s for s in students_all if s["trainee_id"] == str(trainee_id)), None)
    if not student:
        return "Trainee not found", 404

    assignments = load_assignments()
    rem = remaining_for_student(student, students_all, assignments)

    assigned = assignments.get(str(trainee_id))
    return render_template(
        "dashboard.html",
        student=student,
        available=rem["available"],
        total_slots=rem["total_slots"],
        total_remaining=rem["total_remaining"],
        assigned=assigned
    )


@app.route("/assign/<trainee_id>", methods=["POST"])
def assign(trainee_id):
    students_all = load_students()
    student = next((s for s in students_all if s["trainee_id"] == str(trainee_id)), None)
    if not student:
        flash("المتدرب غير موجود.", "error")
        return redirect(url_for("index"))

    assignments = load_assignments()
    chosen = (request.form.get("entity") or "").strip()

    if not chosen:
        flash("اختر جهة تدريب.", "error")
        return redirect(url_for("dashboard", trainee_id=trainee_id))

    rem = remaining_for_student(student, students_all, assignments)
    allowed = {x["name"] for x in rem["available"]}
    if chosen not in allowed:
        flash("هذه الجهة لا يوجد فيها فرص متبقية.", "error")
        return redirect(url_for("dashboard", trainee_id=trainee_id))

    assignments[str(trainee_id)] = {
        "entity": chosen,
        "assigned_at": datetime.now().isoformat(timespec="seconds")
    }
    safe_save_json(ASSIGNMENTS_JSON, assignments)

    flash("تم تأكيد الاختيار وتنقيص الفرص تلقائيًا.", "ok")
    return redirect(url_for("dashboard", trainee_id=trainee_id))


@app.route("/auto_assign/<trainee_id>", methods=["POST"])
def auto_assign(trainee_id):
    students_all = load_students()
    student = next((s for s in students_all if s["trainee_id"] == str(trainee_id)), None)
    if not student:
        flash("المتدرب غير موجود.", "error")
        return redirect(url_for("index"))

    assignments = load_assignments()
    rem = remaining_for_student(student, students_all, assignments)

    if not rem["available"]:
        flash("لا توجد جهات متاحة لهذا التخصص.", "error")
        return redirect(url_for("dashboard", trainee_id=trainee_id))

    chosen = rem["available"][0]["name"]
    assignments[str(trainee_id)] = {
        "entity": chosen,
        "assigned_at": datetime.now().isoformat(timespec="seconds")
    }
    safe_save_json(ASSIGNMENTS_JSON, assignments)

    flash(f"تم الاختيار التلقائي: {chosen}", "ok")
    return redirect(url_for("dashboard", trainee_id=trainee_id))


@app.route("/download_letter/<trainee_id>", methods=["GET"])
def download_letter(trainee_id):
    students_all = load_students()
    student = next((s for s in students_all if s["trainee_id"] == str(trainee_id)), None)
    if not student:
        return "Trainee not found", 404

    assignments = load_assignments()
    assigned = assignments.get(str(trainee_id))
    if not assigned or not assigned.get("entity"):
        flash("لا يمكن الطباعة قبل اختيار جهة التدريب.", "error")
        return redirect(url_for("dashboard", trainee_id=trainee_id))

    entity_name = assigned["entity"]
    try:
        out_path = fill_letter_docx(student, entity_name)
    except Exception as e:
        return f"Error generating letter: {e}", 500

    return send_file(out_path, as_attachment=True, download_name="خطاب_التوجيه.docx")


@app.route("/health")
def health():
    return {"ok": True}


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", "5000")), debug=True)
